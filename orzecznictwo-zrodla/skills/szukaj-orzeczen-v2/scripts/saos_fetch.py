#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
saos_fetch.py - Pobieranie pelnych tresci orzeczen z SAOS Browse API.
Wejscie: plik JSON z szukaj_orzeczen.py lub lista ID.
Wyjscie: JSON + DOCX z pelnymi tresciami.

Uzycie:
    python3 saos_fetch.py --input saos_search_*.json [opcje]
    python3 saos_fetch.py --ids 12345,67890 [opcje]
"""

import argparse, json, os, re, sys, time, urllib.request
from datetime import datetime

SAOS_JUDGMENT_URL = "https://www.saos.org.pl/api/judgments/{id}"
REQUEST_DELAY = 0.5; RETRY_MAX = 3; TIMEOUT = 30

COURT_TYPE_PL = {
    "COMMON": "Sad powszechny", "SUPREME": "Sad Najwyzszy",
    "ADMINISTRATIVE": "Sad administracyjny",
    "CONSTITUTIONAL_TRIBUNAL": "Trybunal Konstytucyjny",
    "NATIONAL_APPEAL_CHAMBER": "Krajowa Izba Odwolawcza",
}
JUDGMENT_TYPE_PL = {
    "SENTENCE": "Wyrok", "DECISION": "Postanowienie",
    "RESOLUTION": "Uchwala", "REGULATION": "Zarzadzenie",
    "REASONS": "Uzasadnienie",
}

def strip_html(text):
    return re.sub(r'<[^>]+>', '', text) if text else ""

def safe_get(d, *keys, default=""):
    current = d
    for key in keys:
        if current is None: return default
        if isinstance(current, dict): current = current.get(key)
        else: return default
    return current if current is not None else default

def sanitize_filename(text, max_len=50):
    clean = re.sub(r'[^\w\s-]', '', text.lower())
    return re.sub(r'[\s]+', '_', clean.strip())[:max_len]

def api_request(url, retries=RETRY_MAX):
    headers = {"Accept": "application/json", "User-Agent": "szukaj-orzeczen/2.0"}
    for attempt in range(retries):
        try:
            req = urllib.request.Request(url, headers=headers)
            resp = urllib.request.urlopen(req, timeout=TIMEOUT)
            return json.loads(resp.read().decode('utf-8'))
        except urllib.error.HTTPError as e:
            if e.code == 429:
                wait = 5 * (attempt + 1)
                print(f"  Rate limit. Czekam {wait}s... ({attempt+1}/{retries})")
                time.sleep(wait)
            elif e.code >= 500:
                wait = 2 * (attempt + 1)
                print(f"  Serwer ({e.code}). Czekam {wait}s...")
                time.sleep(wait)
            else:
                print(f"  HTTP {e.code}"); return None
        except Exception as e:
            wait = 2 * (attempt + 1)
            print(f"  {e}. Czekam {wait}s..."); time.sleep(wait)
    return None

def fetch_judgments(items):
    results = []
    total = len(items)
    print(f"\nPobieram pelne tresci {total} orzeczen...")
    for i, item in enumerate(items):
        jid = item.get("id")
        if not jid: continue
        href = item.get("href", "")
        url = href if href.startswith("http") else SAOS_JUDGMENT_URL.format(id=jid)
        cases = ", ".join(c.get("caseNumber","?") for c in (item.get("courtCases") or []))
        print(f"   [{i+1}/{total}] {cases or jid}...", end=" ")
        data = api_request(url)
        if data:
            judgment = data.get("data") or data
            text_len = len(strip_html(judgment.get("textContent","")))
            print(f"OK ({text_len} zn.)")
            results.append(judgment)
        else:
            print("POMINIETO")
            results.append({"id": jid, "courtCases": item.get("courtCases",[]),
                            "textContent": "", "_fetchError": True})
        time.sleep(REQUEST_DELAY)
    ok = sum(1 for r in results if not r.get("_fetchError"))
    print(f"\nPobrano tresci: {ok}/{total}")
    return results

def generate_judgments_docx(judgments, phrase, timestamp, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn
    except ImportError:
        os.system(f"{sys.executable} -m pip install python-docx --break-system-packages -q")
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.oxml.ns import qn

    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Aptos'; style.font.size = Pt(10)
    title = doc.add_heading('Pelne tresci orzeczen SAOS', level=1)
    for run in title.runs: run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    meta = doc.add_paragraph()
    meta.add_run('Fraza: ').bold = True; meta.add_run(f'\u201e{phrase}\u201d\n')
    meta.add_run('Orzeczen: ').bold = True; meta.add_run(f'{len(judgments)}\n')
    meta.add_run('Data: ').bold = True; meta.add_run(timestamp)

    for idx, j in enumerate(judgments):
        doc.add_page_break()
        cases = j.get("courtCases") or []
        cn = ", ".join(c.get("caseNumber","?") for c in cases)
        jid = j.get("id","?")
        heading = f'{idx+1}. {cn}' if cn else f'{idx+1}. ID: {jid}'
        h = doc.add_heading(heading, level=1)
        for run in h.runs: run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

        court = safe_get(j, "division", "court", "name", default="")
        div_name = safe_get(j, "division", "name", default="")
        ct = COURT_TYPE_PL.get(j.get("courtType",""), j.get("courtType",""))
        jt = JUDGMENT_TYPE_PL.get(j.get("judgmentType",""), j.get("judgmentType",""))
        jd = j.get("judgmentDate","")

        card = doc.add_paragraph()
        card.add_run('Sad: ').bold = True
        card.add_run(f'{court or ct}')
        if div_name: card.add_run(f' \u2014 {div_name}')
        card.add_run('\n')
        card.add_run('Typ: ').bold = True; card.add_run(f'{jt}\n')
        card.add_run('Data: ').bold = True; card.add_run(f'{jd}\n')

        judges = j.get("judges") or []
        if judges:
            card.add_run('Sklad: ').bold = True
            names = []
            for jg in judges:
                n = jg.get("name","")
                roles = jg.get("specialRoles") or []
                if "PRESIDING_JUDGE" in roles: n += " (przew.)"
                if "REPORTING_JUDGE" in roles: n += " (spr.)"
                names.append(n)
            card.add_run(', '.join(names) + '\n')

        kw = j.get("keywords") or []
        if kw:
            card.add_run('Hasla: ').bold = True
            card.add_run(', '.join(kw) + '\n')

        regs = j.get("referencedRegulations") or []
        if regs:
            card.add_run('Przepisy: ').bold = True
            reg_texts = [r.get("journalTitle","") for r in regs[:10]]
            card.add_run('; '.join(filter(None, reg_texts)))
            if len(regs) > 10: card.add_run(f' [+{len(regs)-10}]')
            card.add_run('\n')

        decision = j.get("decision","")
        if decision:
            doc.add_heading('Rozstrzygniecie', level=2)
            doc.add_paragraph(strip_html(decision))

        summary = j.get("summary","")
        if summary:
            doc.add_heading('Teza', level=2)
            doc.add_paragraph(strip_html(summary))

        text = strip_html(j.get("textContent",""))
        if text and not j.get("_fetchError"):
            doc.add_heading('Tresc orzeczenia', level=2)
            chunk = 2000
            for start in range(0, len(text), chunk):
                p = doc.add_paragraph(text[start:start+chunk])
                for run in p.runs: run.font.size = Pt(9)
        else:
            doc.add_paragraph('[Brak tresci orzeczenia w API SAOS]')

    doc.save(output_path)
    print(f"   DOCX: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="Pobieranie pelnych tresci orzeczen SAOS")
    group = parser.add_mutually_exclusive_group(required=True)
    group.add_argument("--input", help="Plik JSON z szukaj_orzeczen.py")
    group.add_argument("--ids", help="Lista ID orzeczen (przecinki)")
    parser.add_argument("--output-dir", default="./saos-output")
    parser.add_argument("--skip-docx", action="store_true")
    args = parser.parse_args()
    os.makedirs(args.output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    print("="*60 + "\n  SZUKAJ ORZECZEN 2.0 - FETCH\n" + "="*60)
    if args.input:
        with open(args.input, 'r', encoding='utf-8') as f: search_data = json.load(f)
        items = search_data.get("items", [])
        phrase = search_data.get("phrase", "unknown")
        print(f"Wczytano {len(items)} orzeczen z: {args.input}")
    else:
        ids = [int(x.strip()) for x in args.ids.split(",") if x.strip()]
        items = [{"id": i, "href": SAOS_JUDGMENT_URL.format(id=i)} for i in ids]
        phrase = f"ids_{len(ids)}"
        print(f"Podano {len(ids)} ID orzeczen")
    if not items: print("Brak orzeczen do pobrania."); sys.exit(0)
    safe = sanitize_filename(phrase)
    judgments = fetch_judgments(items)
    json_path = os.path.join(args.output_dir, f"saos_judgments_{safe}_{ts}.json")
    out = {"phrase": phrase, "count": len(judgments), "timestamp": ts, "judgments": judgments}
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    print(f"   JSON: {json_path}")
    if not args.skip_docx:
        docx_path = os.path.join(args.output_dir, f"saos_judgments_{safe}_{ts}.docx")
        generate_judgments_docx(judgments, phrase, ts, docx_path)
    print(f"\n{'='*60}\n  GOTOWE - {len(judgments)} orzeczen\n{'='*60}")

if __name__ == "__main__":
    main()
