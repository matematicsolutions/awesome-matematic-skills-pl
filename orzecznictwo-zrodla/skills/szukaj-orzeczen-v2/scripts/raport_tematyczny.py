#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
raport_tematyczny.py - Generator raportu tematycznego DOCX.
Wejscie: plik JSON ze struktura raportu (przygotowany przez Claude).
Wyjscie: profesjonalny DOCX z grupowaniem tematycznym.

Uzycie:
    python3 raport_tematyczny.py --input raport_data.json --output raport.docx
"""

import argparse, json, os, sys
from datetime import datetime

FONT_NAME = "Aptos"
HEADER_HEX = "1F4E79"
HEADER_RGB = (0x1F, 0x4E, 0x79)
ALT_HEX = "EAF2FA"
WHITE_RGB = (0xFF, 0xFF, 0xFF)

def ensure_docx():
    try:
        import docx; return True
    except ImportError:
        os.system(f"{sys.executable} -m pip install python-docx --break-system-packages -q")
        return True

def shade_cell(cell, color_hex):
    from docx.oxml.ns import qn
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn('w:shd'),
        {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex})
    tc_pr.append(shd)

def styled_heading(doc, text, level=1):
    from docx.shared import RGBColor
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*HEADER_RGB)
    return h

def add_styled_table(doc, headers, rows):
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h_text in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = h_text
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(*WHITE_RGB)
            run.font.size = Pt(9)
            run.font.name = FONT_NAME
        shade_cell(cell, HEADER_HEX)
    for row_idx, row_data in enumerate(rows):
        row = table.add_row()
        for col_idx, val in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(val) if val is not None else ""
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)
                run.font.name = FONT_NAME
            if row_idx % 2 == 0:
                shade_cell(cell, ALT_HEX)
    doc.add_paragraph()
    return table

def add_title_page(doc, meta):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    for _ in range(4): doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("RAPORT TEMATYCZNY")
    run.bold = True; run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(*HEADER_RGB); run.font.name = FONT_NAME

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p2.add_run("Grupowanie orzecze\u0144 s\u0105dowych")
    run.font.size = Pt(16); run.font.color.rgb = RGBColor(*HEADER_RGB)
    run.font.name = FONT_NAME

    subtitle = meta.get("subtitle", "")
    if subtitle:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p3.add_run(subtitle)
        run.font.size = Pt(13); run.font.color.rgb = RGBColor(*HEADER_RGB)
        run.font.name = FONT_NAME

    for _ in range(3): doc.add_paragraph()

    phrase = meta.get("phrase", "")
    total = meta.get("total_judgments", 0)
    groups = meta.get("total_groups", 0)
    date_range = meta.get("date_range", "")
    report_date = meta.get("report_date", datetime.now().strftime("%d %B %Y"))

    lines = [
        f'Fraza wyszukiwania: \u201e{phrase}\u201d',
        f'Liczba orzecze\u0144: {total}  |  Liczba grup: {groups}',
        f'Zakres dat: {date_range}  |  \u0179r\u00f3d\u0142o: SAOS (www.saos.org.pl)',
        f'Data raportu: {report_date}',
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11); run.font.name = FONT_NAME

    doc.add_paragraph()
    p_d = doc.add_paragraph()
    p_d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_d.add_run(
        "Raport wygenerowany automatycznie na podstawie danych z Systemu Analizy "
        "Orzecze\u0144 S\u0105dowych. Klasyfikacja tematyczna opiera si\u0119 na sygnaturach akt, "
        "has\u0142ach tematycznych, podstawach prawnych, wydziale s\u0105du oraz analizie "
        "tekstowej tre\u015bci orzecze\u0144 (text mining). Klasyfikacja ma charakter przybli\u017cony."
    )
    run.font.size = Pt(9); run.font.name = FONT_NAME; run.italic = True

def add_summary_table(doc, groups, summary_text=""):
    styled_heading(doc, "CZ\u0118\u015a\u0106 I \u2014 PODSUMOWANIE")
    headers = ["Lp.", "Grupa tematyczna", "Orzecze\u0144", "Data od", "Data do", "% zbioru"]
    rows = []
    for i, g in enumerate(groups):
        rows.append([str(i+1), g.get("title",""), str(g.get("count",0)),
                      g.get("date_from","\u2014"), g.get("date_to","\u2014"), g.get("percent","")])
    add_styled_table(doc, headers, rows)
    if summary_text:
        doc.add_paragraph()
        p = doc.add_paragraph(summary_text)
        for run in p.runs: run.font.name = FONT_NAME

def add_group_section(doc, group, group_number):
    title = group.get("title", f"Grupa {group_number}")
    count = group.get("count", 0)
    pct = group.get("percent", "?%")
    desc = group.get("description", "")

    styled_heading(doc, f"CZ\u0118\u015a\u0106 II \u2014 {title.upper()}")
    p = doc.add_paragraph()
    run = p.add_run(f"Liczba orzecze\u0144 w grupie: {count} ({pct} zbioru)")
    run.font.name = FONT_NAME; run.bold = True
    if desc:
        doc.add_paragraph()
        pd = doc.add_paragraph(desc)
        for run in pd.runs: run.font.name = FONT_NAME

    # Patterns
    patterns = group.get("patterns") or []
    if patterns:
        styled_heading(doc, "Wzorce kontekstowe (text mining)", level=3)
        headers = ["Wzorzec", "Wyst\u0105pie\u0144", "% grupy", "% zbioru"]
        rows = [[p.get("pattern",""), str(p.get("count",0)),
                  p.get("pct_group",""), p.get("pct_total","")] for p in patterns]
        add_styled_table(doc, headers, rows)

    # Judgments
    judgments = group.get("judgments") or []
    if judgments:
        styled_heading(doc, "Wykaz orzecze\u0144", level=3)
        headers = ["Lp.", "Sygnatura", "Data", "Typ", "S\u0105d / Wydzia\u0142", "Has\u0142a / Podstawy"]
        rows = [[str(j.get("lp","")), j.get("case_number",""), j.get("date",""),
                  j.get("type",""), j.get("court",""), j.get("keywords_bases","\u2014")]
                 for j in judgments]
        add_styled_table(doc, headers, rows)

    # Legal acts
    acts = group.get("legal_acts") or []
    if acts:
        styled_heading(doc, "Najcz\u0119\u015bciej powo\u0142ywane akty prawne", level=3)
        headers = ["Akt prawny", "Wyst\u0105pie\u0144", "% grupy"]
        rows = [[a.get("act",""), str(a.get("count",0)), a.get("pct","")] for a in acts]
        add_styled_table(doc, headers, rows)

    # Courts
    courts = group.get("courts") or []
    if courts:
        styled_heading(doc, "Rozk\u0142ad s\u0105d\u00f3w", level=3)
        headers = ["S\u0105d", "Orzecze\u0144", "% grupy"]
        rows = [[c.get("court",""), str(c.get("count",0)), c.get("pct","")] for c in courts]
        add_styled_table(doc, headers, rows)

def add_cross_patterns(doc, cross):
    styled_heading(doc, "CZ\u0118\u015a\u0106 III \u2014 WZORCE PRZEKROJOWE")

    gl_acts = cross.get("global_legal_acts") or []
    if gl_acts:
        styled_heading(doc, "3.1 Najcz\u0119\u015bciej powo\u0142ywane akty prawne (globalnie)", level=2)
        headers = ["Akt prawny", "Wyst\u0105pie\u0144", "% zbioru"]
        rows = [[a.get("act",""), str(a.get("count",0)), a.get("pct","")] for a in gl_acts]
        add_styled_table(doc, headers, rows)

    judges = cross.get("top_judges") or []
    if judges:
        styled_heading(doc, "3.2 S\u0119dziowie najcz\u0119\u015bciej orzekaj\u0105cy", level=2)
        headers = ["S\u0119dzia", "Orzecze\u0144", "% zbioru"]
        rows = [[j.get("judge",""), str(j.get("count",0)), j.get("pct","")] for j in judges]
        add_styled_table(doc, headers, rows)

    gl_courts = cross.get("global_courts") or []
    if gl_courts:
        styled_heading(doc, "3.3 Rozk\u0142ad s\u0105d\u00f3w (globalnie)", level=2)
        headers = ["S\u0105d", "Orzecze\u0144", "% zbioru"]
        rows = [[c.get("court",""), str(c.get("count",0)), c.get("pct","")] for c in gl_courts]
        add_styled_table(doc, headers, rows)

    contexts = cross.get("search_contexts", "")
    if contexts:
        styled_heading(doc, "3.4 Konteksty u\u017cycia frazy wyszukiwania", level=2)
        if isinstance(contexts, list):
            for ctx in contexts:
                p = doc.add_paragraph(ctx)
                for run in p.runs: run.font.name = FONT_NAME
        else:
            p = doc.add_paragraph(contexts)
            for run in p.runs: run.font.name = FONT_NAME

    conclusions = cross.get("conclusions") or []
    if conclusions:
        styled_heading(doc, "3.5 Wnioski i rekomendacje", level=2)
        for c in conclusions:
            p = doc.add_paragraph(c)
            for run in p.runs: run.font.name = FONT_NAME

def add_disclaimer(doc, text=None):
    styled_heading(doc, "ZASTRZE\u017bENIA", level=2)
    default = (
        "Niniejszy raport ma charakter wy\u0142\u0105cznie analityczny i nie stanowi porady prawnej. "
        "Dane pochodz\u0105 z publicznego API Systemu Analizy Orzecze\u0144 S\u0105dowych "
        "(SAOS, www.saos.org.pl). Klasyfikacja tematyczna jest przybli\u017cona \u2014 opiera si\u0119 "
        "na sygnaturach akt, has\u0142ach tematycznych, podstawach prawnych, przypisaniu do "
        "wydzia\u0142u s\u0105du oraz analizie tekstowej (text mining) tre\u015bci orzecze\u0144. Pojedyncze "
        "orzeczenia mog\u0105 pasowa\u0107 do wi\u0119cej ni\u017c jednej grupy lub zosta\u0107 sklasyfikowane "
        "nieprawid\u0142owo. Statystyki procentowe odnosz\u0105 si\u0119 wy\u0142\u0105cznie do pobranego zbioru, "
        "nie do ca\u0142o\u015bci orzecznictwa. Pe\u0142ne tre\u015bci orzecze\u0144 dost\u0119pne s\u0105 w bazie SAOS."
    )
    from docx.shared import Pt
    p = doc.add_paragraph(text or default)
    for run in p.runs:
        run.font.name = FONT_NAME; run.font.size = Pt(9); run.italic = True

def generate_report(data, output_path):
    ensure_docx()
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = FONT_NAME; style.font.size = Pt(10)

    meta = data.get("meta", {})
    groups = data.get("groups", [])
    cross = data.get("cross_patterns", {})
    summary_text = data.get("summary", "")
    disclaimer = data.get("disclaimer", None)

    add_title_page(doc, meta)
    doc.add_page_break()

    add_summary_table(doc, groups, summary_text)
    doc.add_page_break()

    for i, group in enumerate(groups):
        add_group_section(doc, group, i+1)
        if i < len(groups) - 1: doc.add_page_break()

    doc.add_page_break()
    if cross: add_cross_patterns(doc, cross)
    doc.add_page_break()
    add_disclaimer(doc, disclaimer)

    doc.save(output_path)
    print(f"Raport zapisany: {output_path}")
    return output_path

def main():
    parser = argparse.ArgumentParser(description="Generator raportu tematycznego DOCX")
    parser.add_argument("--input", required=True, help="JSON ze struktura raportu")
    parser.add_argument("--output", default=None, help="Sciezka DOCX wyjsciowego")
    args = parser.parse_args()
    with open(args.input, 'r', encoding='utf-8') as f: data = json.load(f)
    if not args.output:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        phrase = data.get("meta",{}).get("phrase","raport")
        safe = "".join(c if c.isalnum() or c in '_- ' else '' for c in phrase.lower()).replace(' ','_')[:40]
        args.output = f"raport_tematyczny_{safe}_{ts}.docx"
    print("="*60 + "\n  SZUKAJ ORZECZEN 2.0 - RAPORT TEMATYCZNY\n" + "="*60)
    generate_report(data, args.output)
    print(f"{'='*60}\n  GOTOWE\n{'='*60}")

if __name__ == "__main__":
    main()
