#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
szukaj_orzeczen.py - Przeszukiwanie orzeczeń SAOS API (Search endpoint).
Generuje JSON z metadanymi i DOCX z listą wyników.

Użycie:
    python3 szukaj_orzeczen.py "fraza" [opcje]
"""

import argparse
import json
import os
import re
import sys
import time
import urllib.parse
import urllib.request
from datetime import datetime

SAOS_SEARCH_URL = "https://www.saos.org.pl/api/search/judgments"
PAGE_SIZE = 100
REQUEST_DELAY = 0.5
RETRY_MAX = 3
TIMEOUT = 30

COURT_TYPE_PL = {
    "COMMON": "Sąd powszechny", "SUPREME": "Sąd Najwyższy",
    "ADMINISTRATIVE": "Sąd administracyjny",
    "CONSTITUTIONAL_TRIBUNAL": "Trybunał Konstytucyjny",
    "NATIONAL_APPEAL_CHAMBER": "Krajowa Izba Odwoławcza",
}
JUDGMENT_TYPE_PL = {
    "SENTENCE": "Wyrok", "DECISION": "Postanowienie",
    "RESOLUTION": "Uchwała", "REGULATION": "Zarządzenie",
    "REASONS": "Uzasadnienie",
}

def sanitize_filename(text, max_len=50):
    clean = re.sub(r'[^\w\s-]', '', text.lower())
    clean = re.sub(r'[\s]+', '_', clean.strip())
    return clean[:max_len]

def strip_html(text):
    return re.sub(r'<[^>]+>', '', text) if text else ""

def safe_get(d, *keys, default=""):
    current = d
    for key in keys:
        if current is None: return default
        if isinstance(current, dict): current = current.get(key)
        else: return default
    return current if current is not None else default

def api_request(url, retries=RETRY_MAX):
    headers = {"Accept": "application/json", "User-Agent": "szukaj-orzeczen/2.0"}
    for attempt in range(retries):
        try:
            req = urllib.request.Request(url, headers=headers)
            resp = urllib.request.urlopen(req, timeout=TIMEOUT)
            return json.loads(resp.read().decode('utf-8'))
        except urllib.error.HTTPError as e:
            if e.code == 429:
                wait = 5 * (attempt + 1)
                print(f"  Rate limit (429). Czekam {wait}s... ({attempt+1}/{retries})")
                time.sleep(wait)
            elif e.code >= 500:
                wait = 2 * (attempt + 1)
                print(f"  Blad serwera ({e.code}). Czekam {wait}s... ({attempt+1}/{retries})")
                time.sleep(wait)
            else:
                print(f"  HTTP {e.code}: {e.reason}")
                return None
        except urllib.error.URLError as e:
            wait = 2 * (attempt + 1)
            print(f"  Polaczenie: {e.reason}. Czekam {wait}s... ({attempt+1}/{retries})")
            time.sleep(wait)
        except Exception as e:
            print(f"  Blad: {e}")
            return None
    return None

def build_search_url(phrase, mode, page, sort_field, sort_dir, date_from, date_to):
    params = {"pageSize": PAGE_SIZE, "pageNumber": page,
              "sortingField": sort_field, "sortingDirection": sort_dir}
    if mode == "keywords":
        keywords = [kw.strip() for kw in phrase.split(",") if kw.strip()]
        query_parts = [f"keywords={urllib.parse.quote(kw)}" for kw in keywords]
        base = urllib.parse.urlencode(params)
        url = f"{SAOS_SEARCH_URL}?{base}&{'&'.join(query_parts)}"
    else:
        params["all"] = phrase
        url = f"{SAOS_SEARCH_URL}?{urllib.parse.urlencode(params)}"
    if date_from: url += f"&judgmentDateFrom={date_from}"
    if date_to: url += f"&judgmentDateTo={date_to}"
    return url

def search_judgments(phrase, mode, max_results, sort_field, sort_dir, date_from, date_to):
    all_items, page, total = [], 0, None
    while True:
        url = build_search_url(phrase, mode, page, sort_field, sort_dir, date_from, date_to)
        if page == 0: print(f"Wyszukuje: \"{phrase}\" (tryb: {mode})")
        data = api_request(url)
        if data is None: break
        info = data.get("info") or {}
        if total is None:
            total = info.get("totalResults", 0)
            print(f"   Znaleziono: {total} orzeczen")
        items = data.get("items") or []
        if not items: break
        all_items.extend(items)
        if max_results > 0 and len(all_items) >= max_results:
            all_items = all_items[:max_results]; break
        if len(all_items) >= total: break
        page += 1; time.sleep(REQUEST_DELAY)
    return all_items, total or 0

def generate_search_docx(items, phrase, mode, total_results, timestamp, output_path):
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        os.system(f"{sys.executable} -m pip install python-docx --break-system-packages -q")
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn

    def shade_cell(cell, color_hex):
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = tc_pr.makeelement(qn('w:shd'),
            {qn('w:val'): 'clear', qn('w:color'): 'auto', qn('w:fill'): color_hex})
        tc_pr.append(shd)

    doc = Document()
    style = doc.styles['Normal']; style.font.name = 'Aptos'; style.font.size = Pt(10)
    title = doc.add_heading('Wyniki wyszukiwania SAOS', level=1)
    for run in title.runs: run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    meta = doc.add_paragraph()
    meta.add_run('Fraza: ').bold = True
    meta.add_run(f'\u201e{phrase}\u201d\n')
    meta.add_run('Tryb: ').bold = True; meta.add_run(f'{mode}\n')
    meta.add_run('Znaleziono: ').bold = True; meta.add_run(f'{total_results}  |  ')
    meta.add_run('Pobrano: ').bold = True; meta.add_run(f'{len(items)}\n')
    meta.add_run('Data: ').bold = True; meta.add_run(timestamp)
    doc.add_paragraph()

    headers = ["Lp.", "Sygnatura", "Data", "Typ", "Sad / Wydzial", "Hasla"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True; run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF); run.font.size = Pt(9)
        shade_cell(cell, '1F4E79')

    for idx, item in enumerate(items):
        row = table.add_row()
        cases = ", ".join(c.get("caseNumber","?") for c in (item.get("courtCases") or []))
        court = safe_get(item, "division", "court", "name", default="")
        div_name = safe_get(item, "division", "name", default="")
        court_disp = f"{court} / {div_name}" if court and div_name else court or COURT_TYPE_PL.get(item.get("courtType",""),"")
        jtype = JUDGMENT_TYPE_PL.get(item.get("judgmentType",""), item.get("judgmentType",""))
        kws = ", ".join(item.get("keywords") or [])
        vals = [str(idx+1), cases, item.get("judgmentDate",""), jtype, court_disp, kws]
        for i, val in enumerate(vals):
            cell = row.cells[i]; cell.text = val
            for run in cell.paragraphs[0].runs: run.font.size = Pt(8)
            if idx % 2 == 1: shade_cell(cell, 'EAF2FA')

    doc.save(output_path)
    print(f"   DOCX: {output_path}")

def save_search_json(items, phrase, mode, total_results, timestamp, output_path):
    data = {"phrase": phrase, "mode": mode, "totalResults": total_results,
            "fetchedCount": len(items), "timestamp": timestamp, "items": items}
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    print(f"   JSON: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="Przeszukiwanie orzeczen SAOS")
    parser.add_argument("phrase", help="Fraza wyszukiwania")
    parser.add_argument("--mode", choices=["all","keywords"], default="all")
    parser.add_argument("--max-results", type=int, default=50)
    parser.add_argument("--output-dir", default="./saos-output")
    parser.add_argument("--sort-field", default="JUDGMENT_DATE")
    parser.add_argument("--sort-dir", default="DESC")
    parser.add_argument("--date-from", default=None)
    parser.add_argument("--date-to", default=None)
    args = parser.parse_args()
    if not args.phrase.strip(): print("Fraza nie moze byc pusta."); sys.exit(1)
    os.makedirs(args.output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe = sanitize_filename(args.phrase)
    print("="*60 + "\n  SZUKAJ ORZECZEN 2.0 - SEARCH\n" + "="*60)
    items, total = search_judgments(args.phrase, args.mode, args.max_results,
                                    args.sort_field, args.sort_dir, args.date_from, args.date_to)
    if not items: print("\nBrak wynikow."); sys.exit(0)
    print(f"\nPobrano {len(items)} / {total}")
    print(f"\nZapis do: {args.output_dir}/")
    json_path = os.path.join(args.output_dir, f"saos_search_{safe}_{ts}.json")
    save_search_json(items, args.phrase, args.mode, total, ts, json_path)
    docx_path = os.path.join(args.output_dir, f"saos_search_{safe}_{ts}.docx")
    generate_search_docx(items, args.phrase, args.mode, total, ts, docx_path)
    print(f"\n{'='*60}\n  GOTOWE - {len(items)} orzeczen\n{'='*60}")

if __name__ == "__main__":
    main()
